import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.conf import settings
from django.utils import timezone
from core.constants import DOCXStatus

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "single", "color": "#0000FF"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if not create
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'left', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if not create
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def generate_docx_for_generated_document(generated_doc, overwrite=False):
    """
    Generates a formatted DOCX file for a GeneratedEvidenceDocument.
    """
    if generated_doc.docx_file and not overwrite:
        return generated_doc.docx_file.path

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 1. Header Table
    header_table = doc.add_table(rows=4, cols=4)
    header_table.style = 'Table Grid'
    
    # Merge cells for "Al Shifa Laboratory"
    header_table.cell(0, 0).merge(header_table.cell(0, 2))
    cell = header_table.cell(0, 0)
    cell.text = "Al Shifa Laboratory\nCircular Road, Jaranwala"
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(14)
    
    # Doc Code & Version
    header_table.cell(0, 3).text = f"Doc Code: {generated_doc.document_code}\nVersion: {generated_doc.version}"
    
    # Doc Title
    header_table.cell(1, 0).merge(header_table.cell(1, 3))
    cell = header_table.cell(1, 0)
    cell.text = generated_doc.title.upper()
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(16)
    
    # Dates and Status
    header_table.cell(2, 0).text = "Effective Date:"
    header_table.cell(2, 1).text = "________________"
    header_table.cell(2, 2).text = "Review Date:"
    header_table.cell(2, 3).text = "________________"
    
    header_table.cell(3, 0).text = "Document Type:"
    header_table.cell(3, 1).text = generated_doc.planned_document.get_document_kind_display()
    header_table.cell(3, 2).text = "Status:"
    header_table.cell(3, 3).text = "Waiting for Approval"

    doc.add_paragraph()

    # 2. PHC Mapping Table
    doc.add_heading('PHC/MSDS MAPPING', level=2)
    mapping_table = doc.add_table(rows=1, cols=2)
    mapping_table.style = 'Table Grid'
    
    mapping_data = [
        ('Functional Area', generated_doc.planned_document.batch.functional_area or "N/A"),
        ('Standard', ", ".join(set([i.standard_no for i in generated_doc.planned_document.indicators.all()]))),
        ('Indicator(s)', ", ".join([i.indicator_no for i in generated_doc.planned_document.indicators.all()])),
        ('Evidence Pack', f"{generated_doc.batch.code} - {generated_doc.batch.name}"),
    ]
    
    for label, value in mapping_data:
        row = mapping_table.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = str(value)
    
    doc.add_paragraph()

    # 3. Content Rendering
    doc.add_heading('DOCUMENT BODY', level=2)
    
    # Simple Markdown to DOCX parsing
    content = generated_doc.content_markdown
    # Remove existing signature/header blocks if any (basic heuristic)
    content = re.sub(r'#+ Laboratory Name:.*?\n---', '', content, flags=re.DOTALL)
    content = re.sub(r'### Approval Section.*', '', content, flags=re.DOTALL)
    content = content.strip()
    
    lines = content.split('\n')
    in_table = False
    table_rows = []
    
    for line in lines:
        line = line.strip()
        if not line:
            if in_table:
                render_table(doc, table_rows)
                in_table = False
                table_rows = []
            continue
            
        if line.startswith('## '):
            doc.add_heading(line[3:], level=3)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=4)
        elif line.startswith('|'):
            in_table = True
            table_rows.append(line)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(line[re.search(r'\d+\. ', line).end():], style='List Number')
        else:
            if in_table:
                table_rows.append(line)
            else:
                p = doc.add_paragraph(line)
                # handle basic bold
                for run in p.runs:
                    if '**' in run.text:
                        # This is a very simple bold parser
                        pass

    if in_table:
        render_table(doc, table_rows)

    doc.add_paragraph()

    # 4. Signature Table
    doc.add_heading('APPROVAL & AUTHORIZATION', level=2)
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = 'Table Grid'
    
    cell_prep = sig_table.cell(0, 0)
    p = cell_prep.add_paragraph()
    p.add_run('Prepared by:').bold = True
    p.add_run('\nDr. Muhammad Munaim Tahir\nLab Manager / In-charge\n\nSignature: ___________________\nDate: ___________________')
    
    cell_appr = sig_table.cell(0, 1)
    p = cell_appr.add_paragraph()
    p.add_run('Reviewed and Approved by:').bold = True
    p.add_run('\nDr. Mubasher Ahmed\nConsultant Pathologist\n\nSignature: ___________________\nDate: ___________________')
    
    doc.add_paragraph()
    
    # 5. Footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = f"Al Shifa Laboratory | {generated_doc.document_code} | v{generated_doc.version} | Controlled Document"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)

    # Save
    timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
    batch_dir = os.path.join(settings.MEDIA_ROOT, 'generated_documents_docx', generated_doc.batch.code, timestamp)
    os.makedirs(batch_dir, exist_ok=True)
    
    safe_title = re.sub(r'[^\w\s-]', '', generated_doc.title).replace(' ', '_')
    file_name = f"{generated_doc.document_code}_{safe_title}.docx"
    file_path = os.path.join(batch_dir, file_name)
    
    doc.save(file_path)
    
    # Update model
    relative_path = os.path.relpath(file_path, settings.MEDIA_ROOT)
    generated_doc.docx_file = relative_path
    generated_doc.docx_generated_at = timezone.now()
    generated_doc.docx_status = DOCXStatus.GENERATED_WAITING_APPROVAL
    generated_doc.save()
    
    return file_path

def render_table(doc, rows):
    if len(rows) < 2:
        return
    
    # Simple pipe table parser
    # | Col 1 | Col 2 |
    # | --- | --- |
    # | Data | Data |
    
    header = [c.strip() for c in rows[0].split('|') if c.strip()]
    data_rows = []
    for r in rows[2:]:
        data_rows.append([c.strip() for c in r.split('|') if c.strip()])
        
    if not header:
        return
    
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        
    for row_data in data_rows:
        if len(row_data) != len(header):
            continue
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
